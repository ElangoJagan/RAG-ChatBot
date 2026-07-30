from pathlib import Path
import logging
from dataclasses import dataclass, field
import hashlib

logger = logging.getLogger(__name__)

@dataclass
class Document:
    doc_id: str
    source: str
    page_content: str
    metadata: dict = field(default_factory=dict)

def _hash_id(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()[:16]


class FileLoader:
    def __init__(self, directory: str):
        self.directory = Path(directory)

    def load(self):
        for file in self.directory.iterdir():
            try:
                if file.suffix == '.txt':
                    text = file.read_text(encoding='utf-8')
                elif file.suffix == '.pdf':
                    text = self._extract_pdf(file)
                
                elif file.suffix == '.docx':
                    text = self._extract_docx(file)
                else:
                    continue

                if not text.strip():
                    logger.warning(f"File {file.name} extracted but content is empty, skipping.")
                    continue

                yield Document(
                    doc_id=_hash_id(file.name + text[:100]),
                    source=f"file:{file.name}",
                    page_content=text,
                    metadata={"file_type": file.suffix.lstrip(".")}
                )

            except Exception as e:
                logger.error(f"Failed to load {file.name}: {e}")
            
    def _extract_pdf(self, path):
        import pypdf
        text_parts = []
        with open(path, 'rb')as f:
            reader = pypdf.PdfReader(f)
            for page in reader.pages:
                text = page.extract_text()
                text_parts.append(text)
        return'\n'.join(text_parts)
    
    def _extract_docx(self,path):
        import docx
        doc_path = Path(path)
        doc = docx.Document(doc_path)
        
        text_list = [para.text for para in doc.paragraphs]
        
        full_text = "\n".join(text_list)
        
        return full_text
    
                
        
